# app/services/file_parser.py
"""
HWP 파일 처리 개선
- parse_any_bytes: HWP → PDF 변환 시 convert_bytes_to_pdf_bytes 우선 사용
- parse_any: HWP 로컬 파일 처리 로직 추가
"""
from __future__ import annotations
from typing import Dict, List, Tuple, Union, Optional
import os, re
from io import BytesIO
from app.services.ocr_service import _norm_easyocr_langs

from PIL import Image
if not hasattr(Image, "ANTIALIAS"):
    try:
        Image.ANTIALIAS = Image.LANCZOS
    except Exception:
        from PIL import Image as _I
        Image.ANTIALIAS = getattr(_I, "BICUBIC", None)

def _pdf_page_count(path: str) -> int:
    """PDF 페이지 수 안전 획득"""
    try:
        import fitz
        with fitz.open(path) as d:
            return d.page_count
    except Exception:
        return 1

def _make_image_placeholder_chunk(page_no: int) -> tuple[str, dict]:
    text = f"[page {page_no}: image or low-text content]"
    meta = {
        "type": "image_page",
        "section": "",
        "page": page_no,
        "pages": [page_no],
        "token_count": len(text.split()),
        "bboxes": {},
    }
    return (text, meta)

# ========== [핵심 수정] parse_any_bytes - HWP 처리 개선 ==========
def parse_any_bytes(name_hint: str, content: bytes) -> dict:
    """
    bytes에서 파일 형식 감지 후 파싱
    [핵심 수정] HWP 파일 처리 시 bytes 기반 변환 우선 시도
    
    반환 형식:
      {
        "kind": "pdf"|"docx"|"plain",
        "ext": ".pdf"|...,
        "pages": [...],     # PDF면 페이지별 텍스트
        "blocks": [...],    # PDF면 레이아웃 블록
        "items": [{"text":..}, ...],  # DOCX/PLAIN 등
      }
    """
    ext = sniff_ext_from_name(name_hint)

    # PDF는 그대로
    if ext == ".pdf":
        pages = parse_pdf_pages_from_bytes(content)
        blocks = parse_pdf_blocks_from_bytes(content)
        return {"kind": "pdf", "ext": ext, "pages": pages, "blocks": blocks}

    # DOCX 직접 파싱
    if ext == ".docx":
        items = parse_docx_from_bytes(content)
        return {"kind": "docx", "ext": ext, "items": items}

    # ========== [핵심 수정] HWP 처리 개선 ==========
    if ext in (".hwpx", ".hwp"):
        # 1순위: bytes 기반 PDF 변환 (convert_bytes_to_pdf_bytes)
        from app.services.pdf_converter import convert_bytes_to_pdf_bytes
        try:
            print(f"[PARSE] Attempting HWP→PDF conversion via bytes: {name_hint}")
            pdf_bytes = convert_bytes_to_pdf_bytes(content, ext)
            
            if pdf_bytes:
                print(f"[PARSE] HWP→PDF bytes conversion successful: {len(pdf_bytes)} bytes")
                pages = parse_pdf_pages_from_bytes(pdf_bytes)
                blocks = parse_pdf_blocks_from_bytes(pdf_bytes)
                return {"kind": "pdf", "ext": ".pdf", "pages": pages, "blocks": blocks}
            else:
                print(f"[PARSE] ⚠️ HWP bytes conversion returned None, trying stream converter...")
        except Exception as e:
            print(f"[PARSE] ⚠️ HWP bytes conversion failed: {e}")
        
        # 2순위: 외부 컨버터 (convert_stream_to_pdf_bytes)
        from app.services.pdf_converter import convert_stream_to_pdf_bytes, ConvertStreamError
        try:
            print(f"[PARSE] Attempting HWP→PDF via DOC_CONVERTER_URL...")
            pdf_bytes = convert_stream_to_pdf_bytes(content, ext)
            
            if pdf_bytes:
                print(f"[PARSE] HWP→PDF stream conversion successful: {len(pdf_bytes)} bytes")
                pages = parse_pdf_pages_from_bytes(pdf_bytes)
                blocks = parse_pdf_blocks_from_bytes(pdf_bytes)
                return {"kind": "pdf", "ext": ".pdf", "pages": pages, "blocks": blocks}
        except ConvertStreamError as e:
            print(f"[PARSE] HWP stream conversion failed: {e}")
        except Exception as e:
            print(f"[PARSE] HWP conversion error: {e}")
        
        # 3순위: 실패 시 평문 반환 (fallback)
        print(f"[PARSE] All HWP conversion attempts failed, returning as plain text")
        return {"kind": "plain", "ext": ext, "items": parse_plaintext_bytes(content)}

    # 기타 파일 형식
    return {"kind": "plain", "ext": ext, "items": parse_plaintext_bytes(content)}

# ========== [핵심 수정] parse_any - HWP 로컬 파일 처리 ==========
def parse_any(path: str) -> List[Tuple[int, str]]:
    """
    파일 확장자에 따라 적절한 파서 선택
    [핵심 수정] HWP 파일 처리 로직 추가
    """
    ext = (os.path.splitext(path)[1] or "").lower()
    direct_docx = os.getenv("RAG_PARSE_DIRECT_DOCX", "1") == "1"
    direct_xlsx = os.getenv("RAG_PARSE_DIRECT_XLSX", "1") == "1"
    allow_convert = os.getenv("RAG_CONVERT_NONPDF_TO_PDF", "1") == "1"

    # PDF 직접 파싱
    if ext == ".pdf":
        return parse_pdf(path, by_page=True)

    # DOCX 직접 파싱
    if ext in (".docx",) and direct_docx:
        return parse_docx_sections(path)

    # XLSX/CSV 직접 파싱
    if ext in (".xlsx", ".xlsm", ".csv") and direct_xlsx:
        if ext == ".csv":
            import csv
            lines = []
            with open(path, "r", encoding="utf-8", errors="ignore", newline="") as f:
                rdr = csv.reader(f)
                for row in rdr:
                    lines.append(" | ".join([c.strip() for c in row]))
            return [(1, "\n".join(lines))]
        return parse_xlsx_tables(path)

    # ========== [핵심 수정] HWP 파일 명시적 처리 ==========
    if ext in (".hwp", ".hwpx"):
        if allow_convert:
            print(f"[PARSE] HWP file detected, converting to PDF: {path}")
            from app.services.pdf_converter import convert_to_pdf, ConvertError
            try:
                pdf_path = convert_to_pdf(path)
                print(f"[PARSE] HWP→PDF conversion successful: {pdf_path}")
                return parse_pdf(pdf_path, by_page=True)
            except ConvertError as e:
                raise RuntimeError(f"HWP 변환 실패: {e}")
        else:
            raise RuntimeError(f"HWP 파일 변환이 비활성화되어 있습니다 (RAG_CONVERT_NONPDF_TO_PDF=0): {path}")

    # 기타 파일 형식 → PDF 변환
    if allow_convert:
        from app.services.pdf_converter import convert_to_pdf
        pdf_path = convert_to_pdf(path)
        return parse_pdf(pdf_path, by_page=True)

    raise RuntimeError(f"Unsupported file type: {ext}")

# ========== 나머지 함수들은 기존 코드 유지 ==========

def sniff_ext_from_name(name: str) -> str:
    return (os.path.splitext(name)[1] or "").lower()

def parse_pdf_pages_from_bytes(pdf_bytes: bytes) -> List[str]:
    """bytes에서 페이지별 텍스트 추출 (pdfminer.six 기반)"""
    import pdfminer.high_level
    from pdfminer.layout import LAParams, LTTextContainer
    from io import BytesIO
    
    laparams = LAParams()
    pages: List[str] = []
    for layout in pdfminer.high_level.extract_pages(BytesIO(pdf_bytes), laparams=laparams):
        parts = []
        for elem in layout:
            if isinstance(elem, LTTextContainer):
                parts.append(elem.get_text())
        text = "".join(parts).strip()
        pages.append(text)
    return pages

def parse_pdf_blocks_from_bytes(pdf_bytes: bytes) -> List[Tuple[int, List[Dict]]]:
    """bytes에서 레이아웃 블록 추출"""
    import fitz
    out = []
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        for i, page in enumerate(doc, start=1):
            blocks = []
            for b in page.get_text("blocks"):
                if not b or len(b) < 5:
                    continue
                x0, y0, x1, y1, txt = b[0], b[1], b[2], b[3], (b[4] or "").strip()
                if txt:
                    blocks.append({"text": txt, "bbox": [float(x0), float(y0), float(x1), float(y1)]})
            out.append((i, blocks))
    finally:
        doc.close()
    return out

def parse_docx_from_bytes(content: bytes) -> List[Dict]:
    """DOCX bytes → 라인 단위"""
    from docx import Document
    from io import BytesIO
    
    doc = Document(BytesIO(content))
    items = []
    for p in doc.paragraphs:
        line = p.text.strip()
        if line:
            items.append({"text": line})
    return items

def parse_plaintext_bytes(content: bytes) -> List[Dict]:
    """평문 bytes → 라인 단위"""
    try:
        text = content.decode("utf-8", errors="ignore")
    except:
        text = ""
    lines = [ln.strip() for ln in text.split("\n") if ln.strip()]
    return [{"text": ln} for ln in lines]

# ========== 기존 파싱 함수들 (변경 없음) ==========

def _extract_text_pdfminer(path: str, by_page: bool) -> Union[str, List[Tuple[int, str]]]:
    import pdfminer.high_level
    from pdfminer.layout import LAParams, LTTextContainer

    laparams = LAParams()
    if not by_page:
        return (pdfminer.high_level.extract_text(path, laparams=laparams) or "").strip()

    pages: List[Tuple[int, str]] = []
    for i, layout in enumerate(pdfminer.high_level.extract_pages(path, laparams=laparams), start=1):
        parts = []
        for elem in layout:
            if isinstance(elem, LTTextContainer):
                parts.append(elem.get_text())
        text = "".join(parts).strip()
        if text:
            pages.append((i, text))
    return pages

def _extract_text_pypdf2(path: str, by_page: bool) -> Union[str, List[Tuple[int, str]]]:
    from PyPDF2 import PdfReader
    reader = PdfReader(path)
    if not by_page:
        txt = "\n\n".join([(p.extract_text() or "") for p in reader.pages]).strip()
        return txt
    pages: List[Tuple[int, str]] = []
    for i, p in enumerate(reader.pages, start=1):
        t = (p.extract_text() or "").strip()
        if t:
            pages.append((i, t))
    return pages

def _render_pdf_pages_fitz(path: str, dpi: int):
    import fitz
    import numpy as np

    zoom = max(1.0, dpi / 72.0)
    mat = fitz.Matrix(zoom, zoom)
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc, start=1):
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
            yield i, img
    finally:
        doc.close()

def _ocr_with_paddle(images_iter, by_page: bool, lang: str) -> Union[str, List[Tuple[int, str]]]:
    from paddleocr import PaddleOCR
    ocr = PaddleOCR(lang=("korean" if lang in ("kor", "korean", "ko") else lang), show_log=False)
    pages: List[Tuple[int, str]] = []
    for pno, img in images_iter:
        result = ocr.ocr(img, cls=True)
        text = " ".join([line[1][0] for line in (result[0] or []) if line[1] and line[1][0]]).strip()
        if text:
            pages.append((pno, text))
    if by_page:
        return pages
    return "\n\n".join([t for _, t in pages]).strip()

def _ocr_with_tesseract(images_iter, by_page: bool, lang: str) -> Union[str, List[Tuple[int, str]]]:
    import pytesseract
    from PIL import Image
    
    tcmd = os.getenv("OCR_TESSERACT_CMD")
    if tcmd:
        pytesseract.pytesseract.tesseract_cmd = tcmd

    pages: List[Tuple[int, str]] = []
    for pno, img in images_iter:
        pil = Image.fromarray(img)
        text = pytesseract.image_to_string(pil, lang=(lang or "kor+eng")).strip()
        if text:
            pages.append((pno, text))
    if by_page:
        return pages
    return "\n\n".join([t for _, t in pages]).strip()

def _ocr_with_easyocr(images_iter, by_page: bool, lang: str) -> Union[str, List[Tuple[int, str]]]:
    import easyocr
    import numpy as np
    from PIL import Image
    
    langs = _norm_easyocr_langs(lang)
    gpu = os.getenv("OCR_EASYOCR_GPU", "1").strip() == "1"
    
    width_ths = float(os.getenv("OCR_WIDTH_THS", "0.5"))
    height_ths = float(os.getenv("OCR_HEIGHT_THS", "0.5"))
    
    reader = easyocr.Reader(langs, gpu=gpu)
    pages: List[Tuple[int, str]] = []
    
    for pno, img in images_iter:
        if img.ndim == 3 and img.shape[2] == 4:
            img = img[:, :, :3]
        
        best_text = ""
        
        for angle in [0, 90, 180, 270]:
            if angle != 0:
                pil_img = Image.fromarray(img)
                pil_img = pil_img.rotate(-angle, expand=True)
                rotated_img = np.array(pil_img)
            else:
                rotated_img = img
            
            res = reader.readtext(
                rotated_img,
                detail=1,
                paragraph=False,
                width_ths=width_ths,
                height_ths=height_ths,
            )
            
            if res:
                sorted_res = sorted(res, key=lambda x: (x[0][0][1], x[0][0][0]))
                text = " ".join([x[1] for x in sorted_res if len(x) >= 2 and x[1]]).strip()
                
                if len(text) > len(best_text):
                    best_text = text
        
        if best_text:
            pages.append((pno, best_text))
    
    if by_page:
        return pages
    return "\n\n".join([t for _, t in pages]).strip()

def _has_cid_codes(text: str) -> bool:
    if not text:
        return False
    
    cid_pattern = r'\(cid:\d+\)'
    cid_matches = re.findall(cid_pattern, text)
    cid_count = len(cid_matches)
    
    if cid_count < 10:
        return False
    
    total_chars = len(text)
    cid_chars = sum(len(m) for m in cid_matches)
    cid_ratio = cid_chars / max(1, total_chars)
    
    is_cid = cid_ratio > 0.05
    
    if is_cid:
        print(f"⚠️ CID codes detected: {cid_count} codes, {cid_ratio*100:.1f}% of text")
    
    return is_cid

def _should_ocr(txt_or_pages: Union[str, List[Tuple[int, str]]]) -> bool:
    try:
        th = int(os.getenv("OCR_MIN_CHARS", "40"))
    except Exception:
        th = 40
    
    if isinstance(txt_or_pages, str):
        if len(txt_or_pages.strip()) < th:
            return True
        if _has_cid_codes(txt_or_pages):
            return True
        return False
    
    total = sum(len(t or "") for _, t in (txt_or_pages or []))
    if total < th:
        return True
    
    for _, text in (txt_or_pages or []):
        if _has_cid_codes(text):
            return True
    
    return False

def parse_docx_sections(path: str) -> List[Tuple[int, str]]:
    from docx import Document
    doc = Document(path)
    sections: List[Tuple[int, str]] = []
    cur = []
    sec_no = 0

    def flush():
        nonlocal cur, sec_no
        if cur:
            sec_no += 1
            sections.append((sec_no, "\n".join(cur).strip()))
            cur.clear()

    for p in doc.paragraphs:
        style = (p.style.name or "").lower() if p.style else ""
        line = p.text.strip()
        if not line:
            continue
        if "heading" in style:
            flush()
            cur.append(line)
        else:
            cur.append(line)

    for t in doc.tables:
        lines = []
        for r in t.rows:
            cells = [c.text.strip().replace("\n", " ") for c in r.cells]
            lines.append(" | ".join(cells))
        if lines:
            cur.append("[표]\n" + "\n".join(lines))

    flush()
    return sections

def parse_xlsx_tables(path: str) -> List[Tuple[int, str]]:
    import openpyxl
    wb = openpyxl.load_workbook(path, data_only=True)
    out: List[Tuple[int, str]] = []
    for si, ws in enumerate(wb.worksheets, start=1):
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            continue
        headers = [str(h or "").strip() for h in rows[0]]
        lines = [f"[시트] {ws.title}"]
        for r in rows[1:]:
            pairs = []
            for h, v in zip(headers, r):
                if h:
                    pairs.append(f"{h}: {'' if v is None else str(v)}")
            if pairs:
                lines.append(", ".join(pairs))
        txt = "\n".join(lines).strip()
        if txt:
            out.append((si, txt))
    return out

def parse_pdf(path: str, by_page: bool = False) -> Union[str, List[Tuple[int, str]]]:
    """
    PDF 파싱 with OCR 지원
    HWP 변환 PDF는 OCR 스킵
    """
    import fitz
    
    skip_ocr = False
    try:
        doc = fitz.open(path)
        metadata = doc.metadata
        if metadata.get('subject') == 'NO_OCR_NEEDED' or metadata.get('title') == 'HWP_CONVERTED':
            skip_ocr = True
            print(f"[PARSE] Detected HWP-converted PDF, skipping OCR")
        doc.close()
    except Exception:
        pass
    
    ocr_mode = os.getenv("OCR_MODE", "auto").lower()
    # HWP 변환 PDF는 강제로 OCR 스킵
    if skip_ocr:
        ocr_mode = "never"
    ocr_engine = os.getenv("OCR_ENGINE", "easyocr").lower()
    ocr_dpi = int(os.getenv("OCR_DPI", "300"))

    page_count = _pdf_page_count(path)

    if ocr_engine == "tesseract":
        ocr_langs = os.getenv("OCR_LANGS", "kor+eng")
    elif ocr_engine == "easyocr":
        ocr_lang = os.getenv("OCR_LANG", "ko,en")
    else:
        ocr_lang = os.getenv("OCR_LANG", "korean")

    text_result: Optional[Union[str, List[Tuple[int, str]]]] = None
    if ocr_mode in ("auto", "never"):
        try:
            text_result = _extract_text_pdfminer(path, by_page)
        except Exception:
            try:
                text_result = _extract_text_pypdf2(path, by_page)
            except Exception:
                text_result = "" if not by_page else []
        if ocr_mode == "never":
            return text_result if text_result is not None else ("" if not by_page else [])
        if text_result and not _should_ocr(text_result):
            return text_result

    try:
        images_iter = _render_pdf_pages_fitz(path, ocr_dpi)

        def _empty(x):
            return (isinstance(x, str) and not x.strip()) or (isinstance(x, list) and len(x) == 0)

        if ocr_engine == "tesseract":
            ocr_out = _ocr_with_tesseract(images_iter, by_page, ocr_langs)
        elif ocr_engine == "easyocr":
            ocr_out = _ocr_with_easyocr(images_iter, by_page, ocr_lang)
        else:
            ocr_out = _ocr_with_paddle(images_iter, by_page, ocr_lang)

        if _empty(ocr_out):
            images_iter = _render_pdf_pages_fitz(path, ocr_dpi)
            if ocr_engine == "easyocr":
                ocr_out = _ocr_with_tesseract(images_iter, by_page, os.getenv("OCR_LANGS", "kor+eng"))
            else:
                ocr_out = _ocr_with_easyocr(images_iter, by_page, os.getenv("OCR_LANG", "ko,en"))

        if _empty(ocr_out):
            if text_result and ((isinstance(text_result, str) and text_result.strip()) or (isinstance(text_result, list) and len(text_result) > 0)):
                return text_result
            if by_page:
                return [(p, _make_image_placeholder_chunk(p)[0]) for p in range(1, page_count + 1)]
            else:
                return "\n".join(_make_image_placeholder_chunk(p)[0] for p in range(1, page_count + 1))

        return ocr_out

    except Exception as e:
        if text_result and ((isinstance(text_result, str) and text_result.strip()) or (isinstance(text_result, list) and len(text_result) > 0)):
            return text_result
        if by_page:
            return [(p, _make_image_placeholder_chunk(p)[0]) for p in range(1, page_count + 1)]
        else:
            return "\n".join(_make_image_placeholder_chunk(p)[0] for p in range(1, page_count + 1))

def parse_pdf_blocks(path: str) -> list[tuple[int, list[dict]]]:
    import fitz
    out = []
    doc = fitz.open(path)
    try:
        for i, page in enumerate(doc, start=1):
            blocks = []
            for b in page.get_text("blocks"):
                if not b or len(b) < 5:
                    continue
                x0, y0, x1, y1, txt = b[0], b[1], b[2], b[3], (b[4] or "").strip()
                if txt:
                    blocks.append({
                        "text": txt,
                        "bbox": {"x0": float(x0), "y0": float(y0), "x1": float(x1), "y1": float(y1)}
                    })
            out.append((i, blocks))
    finally:
        doc.close()
    return out